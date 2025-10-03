# -*- coding: utf-8 -*-
"""
Générateur de documents Word à partir de modèles
"""
import logging
from pathlib import Path
from typing import List, Dict, Any, Tuple
from docx import Document
from docx2pdf import convert

from config import PLACEHOLDER, OUT_DOCX_DIR, OUT_PDF_DIR
from file_utils import safe_filename, safe_email_for_filename


class DocumentGenerator:
    """Classe pour générer des documents Word à partir de modèles."""
    
    def __init__(self, template_path: Path):
        self.template_path = template_path
        if not template_path.exists():
            raise FileNotFoundError(f"Modèle introuvable: {template_path}")
    
    def replace_placeholder_in_paragraph(self, paragraph, placeholder: str, replacement: str) -> None:
        """Remplace un placeholder dans un paragraphe."""
        if placeholder in paragraph.text:
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, replacement)
    
    def force_replace_across_runs(self, paragraph, placeholder: str, replacement: str) -> bool:
        """Force le remplacement d'un placeholder qui s'étend sur plusieurs runs."""
        full = paragraph.text
        if placeholder in full:
            new_text = full.replace(placeholder, replacement)
            # Purger les runs et réécrire
            for _ in range(len(paragraph.runs)-1, -1, -1):
                try:
                    paragraph.runs[_].text = ""
                except Exception:
                    pass
            try:
                paragraph.clear()
            except Exception:
                # Fallback pour les versions plus anciennes
                while paragraph.runs:
                    paragraph.runs[-1].text = ""
            paragraph.add_run(new_text)
            return True
        return False
    
    def replace_in_tables(self, doc, placeholder: str, replacement: str) -> None:
        """Remplace les placeholders dans les tableaux."""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if placeholder in paragraph.text:
                            self.replace_placeholder_in_paragraph(paragraph, placeholder, replacement) or \
                            self.force_replace_across_runs(paragraph, placeholder, replacement)
    
    def generate_document(self, name: str, index: int) -> Path:
        """Génère un document Word pour un nom donné."""
        doc = Document(str(self.template_path))
        
        # Remplacer dans les paragraphes
        for p in doc.paragraphs:
            if PLACEHOLDER in p.text:
                self.replace_placeholder_in_paragraph(p, PLACEHOLDER, name)
                if PLACEHOLDER in p.text:
                    self.force_replace_across_runs(p, PLACEHOLDER, name)
        
        # Remplacer dans les tableaux
        self.replace_in_tables(doc, PLACEHOLDER, name)
        
        # Sauvegarder
        out_path = OUT_DOCX_DIR / f"{safe_filename(name)}.docx"
        doc.save(str(out_path))
        return out_path
    
    def convert_to_pdf(self, docx_path: Path, email: str = "") -> Path:
        """Convertit un document Word en PDF."""
        email_suffix = ("_" + safe_email_for_filename(email)) if email else ""
        pdf_path = OUT_PDF_DIR / (docx_path.stem + f"{email_suffix}.pdf")
        convert(str(docx_path), str(pdf_path))
        return pdf_path
    
    def generate_documents_batch(self, rows: List[Dict[str, Any]], retry_count: int = 3) -> Tuple[List[Path], List[Path]]:
        """Génère tous les documents pour une liste de données avec gestion d'erreurs robuste."""
        docx_files = []
        pdf_files = []
        errors = []

        for i, row in enumerate(rows):
            name = row.get('nom', 'inconnu')
            attempts = 0
            success = False

            while attempts < retry_count and not success:
                try:
                    # Générer le document Word
                    docx_path = self.generate_document(name, i + 1)
                    docx_files.append(docx_path)

                    # Convertir en PDF
                    pdf_path = self.convert_to_pdf(docx_path, row.get('email', ''))
                    pdf_files.append(pdf_path)

                    logging.info(f"Document généré: {docx_path.name} -> {pdf_path.name}")
                    success = True

                except Exception as e:
                    attempts += 1
                    error_msg = f"Erreur lors de la génération du document pour {name} (tentative {attempts}/{retry_count}): {e}"

                    if attempts < retry_count:
                        logging.warning(error_msg)
                    else:
                        logging.error(error_msg)
                        errors.append({
                            'nom': name,
                            'erreur': str(e),
                            'index': i
                        })

        if errors:
            error_summary = "\n".join([f"- {err['nom']}: {err['erreur']}" for err in errors])
            logging.error(f"Échecs de génération ({len(errors)}/{len(rows)}):\n{error_summary}")

            # Ne pas lever d'exception si au moins un document a été généré
            if not docx_files:
                raise Exception(f"Tous les documents ont échoué. Première erreur: {errors[0]['erreur']}")

        return docx_files, pdf_files
