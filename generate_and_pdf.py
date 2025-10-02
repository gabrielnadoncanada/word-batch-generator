# -*- coding: utf-8 -*-
"""
Génère des .docx à partir d'un modèle Word (placeholder {{VENDEUR}}),
convertit en PDF (docx2pdf via Outlook/Word), puis (optionnel) envoie
chaque PDF par e-mail via Outlook Desktop avec la signature système.
- Un PDF par entrée CSV (aucune fusion).
- CSV: colonnes minimales 'nom' et optionnel 'email' (pour nom de fichier et envoi).
"""

import csv
import sys
import os
from pathlib import Path
from typing import List
from docx import Document
from docx2pdf import convert
import logging, traceback, platform, struct
from datetime import datetime


BASE_DIR = Path(__file__).resolve().parent
TEMPLATE = BASE_DIR / "templates" / "modele.docx"
CSV_FILE = BASE_DIR / "data" / "entrepreneurs.csv"
OUT_DOCX_DIR = BASE_DIR / "out" / "docx"
OUT_PDF_DIR  = BASE_DIR / "out" / "pdf"
PLACEHOLDER = "{{VENDEUR}}"
LOG_DIR = OUT_PDF_DIR.parent / "logs"
LOG_DIR.mkdir(parents=True, exist_ok=True)
LOG_FILE = LOG_DIR / "mail.log"

# ==== Envoi e-mail (Outlook Desktop) ====
SEND_EMAIL = True           # Passe à True pour activer l'envoi après génération
FROM_ACCOUNT = "gabriel@dilamco.com"          # None = compte par défaut d'Outlook, sinon adresse explicite
CC = ""                      # ex: "cc1@example.com;cc2@example.com"
BCC = ""                     # ex: "bcc@example.com"
SUBJECT_TEMPLATE = "Soumission - 25142 - École Arc-en-ciel Pavillon 1 (Laval)"
BODY_HTML_TEMPLATE = "Bonjour, <br>Je me permets de vous transmettre notre soumission pour le projet 25142 - École Arc-en-ciel Pavillon 1 (Laval).<br>Vous trouverez ci-joint le document détaillant notre proposition pour la fourniture et l'installation de la vanité.<br>Nous restons disponibles pour toute précision ou information complémentaire."

USE_SYSTEM_SIGNATURE = True  # Ajoute la signature par défaut du système
# --- Choix explicite de signature ---
SIGNATURE_NAME = "Dilamco (gabriel@dilamco.com)"    # <- mets ici le NOM EXACT du fichier .htm de ta signature (sans extension).
                                 #    Laisse "" ou None pour ne pas forcer et utiliser la signature système.

STRICT_EMBED_SIGNATURE_IMAGES = False
# Si True : on attache les images de la signature (dossier "<Nom>_files") et on les inline en CID.
# Utile si ta signature contient des logos/images qui n'apparaissent pas chez certains destinataires.

# =======================================

OUT_DOCX_DIR.mkdir(parents=True, exist_ok=True)
OUT_PDF_DIR.mkdir(parents=True, exist_ok=True)

logging.basicConfig(
    level=logging.DEBUG,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.FileHandler(LOG_FILE, encoding="utf-8"),
        logging.StreamHandler()
    ],
)

def log_env():
    try:
        import win32com, pythoncom
        win32ver = getattr(win32com, "__version__", "?")
    except Exception:
        win32ver = "unavailable"
    logging.info("==== RUNTIME ====")
    logging.info(f"Python: {platform.python_version()} ({platform.architecture()[0]})")
    logging.info(f"OS: {platform.platform()}")
    logging.info(f"Pywin32: {win32ver}")
    logging.info(f"Working dir: {Path.cwd()}")
    logging.info(f"Script dir: {BASE_DIR}")
    logging.info(f"Docx out: {OUT_DOCX_DIR} | Pdf out: {OUT_PDF_DIR} | Log: {LOG_FILE}")

def replace_placeholder_in_paragraph(paragraph, placeholder: str, replacement: str) -> None:
    if placeholder in paragraph.text:
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, replacement)


def force_replace_across_runs(paragraph, placeholder: str, replacement: str) -> bool:
    full = paragraph.text
    if placeholder in full:
        new_text = full.replace(placeholder, replacement)
        # purge runs and rewrite
        for _ in range(len(paragraph.runs)-1, -1, -1):
            try:
                paragraph.runs[_].text = ""
            except Exception:
                pass
        try:
            paragraph.clear()
        except Exception:
            # older python-docx may not expose .clear(); fallback
            while paragraph.runs:
                paragraph.runs[-1].text = ""
        paragraph.add_run(new_text)
        return True
    return False


def replace_in_tables(doc, placeholder: str, replacement: str) -> None:
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if placeholder in paragraph.text:
                        replace_placeholder_in_paragraph(paragraph, placeholder, replacement) or force_replace_across_runs(paragraph, placeholder, replacement)


def safe_filename(name: str) -> str:
    return "".join(c for c in name if c.isalnum() or c in (" ", "-", "_")).strip().replace(" ", "_")


def safe_email_for_filename(email: str) -> str:
    email = (email or "").replace("@", "_at_")
    return "".join(c for c in email if c.isalnum() or c in ("-", "_", "."))

def _read_text_smart(path: Path) -> str:
    data = path.read_bytes()
    for enc in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    # dernier recours : ne PAS perdre les chars (remplacement visible)
    return data.decode("utf-8", errors="replace")


def make_docx_for(name: str, index: int) -> Path:
    doc = Document(str(TEMPLATE))
    for p in doc.paragraphs:
        if PLACEHOLDER in p.text:
            replace_placeholder_in_paragraph(p, PLACEHOLDER, name)
            if PLACEHOLDER in p.text:
                force_replace_across_runs(p, PLACEHOLDER, name)
    replace_in_tables(doc, PLACEHOLDER, name)

    out_path = OUT_DOCX_DIR / f"{index:02d}_{safe_filename(name)}.docx"
    doc.save(str(out_path))
    return out_path


def convert_docx_to_pdf(docx_path: Path, email: str = "") -> Path:
    email_suffix = ("_" + safe_email_for_filename(email)) if email else ""
    pdf_path = OUT_PDF_DIR / (docx_path.stem + f"{email_suffix}.pdf")
    convert(str(docx_path), str(pdf_path))
    return pdf_path


def read_rows(csv_file: Path) -> List[dict]:
    rows: List[dict] = []
    with open(csv_file, newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            name = (row.get("nom") or "").strip()
            email = (row.get("email") or "").strip()
            if name:
                rows.append({"nom": name, "email": email})
    return rows


# ---------- Outlook helpers (signature + send) ----------
def _detect_default_signature_name() -> str:
    try:
        import winreg
    except Exception:
        return ""
    for ver in ("16.0", "15.0", "14.0"):
        try:
            key_path = fr"Software\Microsoft\Office\{ver}\Common\MailSettings"
            with winreg.OpenKey(winreg.HKEY_CURRENT_USER, key_path) as k:
                name, _ = winreg.QueryValueEx(k, "NewSignature")
                if name:
                    return name
        except FileNotFoundError:
            continue
        except Exception:
            continue
    return ""


def _load_system_signature_html() -> str:
    sig_dir = Path(os.environ.get("APPDATA", "")) / "Microsoft" / "Signatures"
    if not sig_dir.exists():
        return ""
    name = _detect_default_signature_name()
    candidate = None
    if name:
        p = sig_dir / f"{name}.htm"
        if p.exists():
            candidate = p
        else:
            p2 = sig_dir / f"{name}.html"
            if p2.exists():
                candidate = p2
    if not candidate:
        files = list(sig_dir.glob("*.htm")) or list(sig_dir.glob("*.html"))
        if files:
            candidate = files[0]
    return _read_text_smart(candidate) if candidate and candidate.exists() else ""


def debug_list_outlook_accounts():
    try:
        import win32com.client as win32
        outlook = win32.gencache.EnsureDispatch("Outlook.Application")
        session = outlook.Session
        logging.info("Comptes Outlook détectés:")
        for acct in session.Accounts:
            logging.info(f" - DisplayName={getattr(acct, 'DisplayName', '')} | SMTP={getattr(acct, 'SmtpAddress', '')}")
    except Exception as e:
        logging.error("Echec debug_list_outlook_accounts()")
        logging.error(traceback.format_exc())

def ensure_mapi_ready():
    """Force une session MAPI opérationnelle (évitant CreateItem/Send refusés)."""
    import pythoncom, time
    import win32com.client as win32
    from win32com.client import constants

    try:
        pythoncom.CoInitialize()
    except Exception:
        pass

    outlook = win32.gencache.EnsureDispatch("Outlook.Application")
    ns = outlook.GetNamespace("MAPI")
    try:
        ns.Logon("", "", False, False)  # profil par défaut si déjà connecté -> no-op
    except Exception:
        logging.warning("ns.Logon a échoué (probablement déjà logué), on continue.")

    # réveille la MAPI / store
    try:
        _ = ns.GetDefaultFolder(6)  # 6=olFolderInbox
    except Exception:
        logging.warning("GetDefaultFolder(Inbox) a échoué.")
    return outlook, ns

def _signatures_dir() -> Path:
    return Path(os.environ.get("APPDATA", "")) / "Microsoft" / "Signatures"

def list_signatures() -> list[str]:
    d = _signatures_dir()
    if not d.exists(): 
        return []
    names = []
    for p in d.glob("*.htm"):
        names.append(p.stem)  # sans extension
    for p in d.glob("*.html"):
        if p.stem not in names:
            names.append(p.stem)
    return sorted(names)

def _load_signature_html_by_name(name: str) -> tuple[str, str]:
    if not name:
        return "", ""
    d = Path(os.environ.get("APPDATA", "")) / "Microsoft" / "Signatures"
    for ext in ("htm", "html"):
        path = d / f"{name}.{ext}"
        if path.exists():
            return _read_text_smart(path), path.stem
    return "", ""


def _attach_and_inline_signature_images(mail, signature_name: str, signature_html: str) -> str:
    """
    Joindre les images de %APPDATA%/Microsoft/Signatures/<name>_files et
    remplacer les src par cid: pour une livraison fiable.
    """
    if not signature_name or not signature_html:
        return signature_html
    base = _signatures_dir() / f"{signature_name}_files"
    if not base.exists():
        return signature_html

    # On attache chaque image et remplace sa source par un CID unique.
    # NB: Outlook acceptera PNG/JPG/GIF ; on conserve le filename comme CID par simplicité.
    import re
    from uuid import uuid4

    def attach_and_replace(match):
        src = match.group(1)  # chemin relatif ex: "Dilamco_FR_files/image001.png"
        # Trouver le fichier réel (src peut déjà être "name_files/xxx" ou juste "xxx")
        candidate = base / Path(src).name
        if not candidate.exists():
            # Essaye le chemin tel quel si l'export contient des sous-dossiers
            candidate = _signatures_dir() / src
            if not candidate.exists():
                return match.group(0)  # laisse tel quel

        try:
            att = mail.Attachments.Add(str(candidate))
            # Définir un CID stable basé sur le nom (évite briser les liens lors de replies)
            cid = Path(candidate).name.replace(".", "_")
            # PR_ATTACH_CONTENT_ID = 0x3712001F
            att.PropertyAccessor.SetProperty("http://schemas.microsoft.com/mapi/proptag/0x3712001F", cid)
            return f'src="cid:{cid}"'
        except Exception:
            return match.group(0)

    # Remplacer src="...".
    # Couvre src="..." et src='...'
    new_html = re.sub(r'src=["\']([^"\']+)["\']', attach_and_replace, signature_html)
    return new_html


def send_email_via_outlook(
    to_email: str,
    subject: str,
    html_body: str,
    attachments: list,
    from_account: str | None = None,
    cc: str = "",
    bcc: str = "",
    use_signature: bool = True,
    max_retries: int = 5,
    delay_seconds: float = 2.0,
) -> None:
    import time, pythoncom
    import win32com.client as win32
    from win32com.client import constants

    logging.info(f"[MAIL] Préparation -> to={to_email} | from={from_account} | subj={subject}")

    try:
        pythoncom.CoInitialize()
    except Exception:
        pass

    try:
        outlook, ns = ensure_mapi_ready()
    except Exception as e:
        logging.error("[MAIL] ensure_mapi_ready failed: %s", e)
        logging.error(traceback.format_exc())
        raise

    # CREATE ITEM avec retries (Outlook peut être occupé)
    mail = None
    for attempt in range(1, max_retries + 1):
        try:
            mail = outlook.CreateItem(constants.olMailItem)  # 0
            logging.debug("[MAIL] CreateItem OK (attempt %d)", attempt)
            break
        except Exception as e:
            msg = str(e)
            logging.warning("[MAIL] CreateItem échoué (attempt %d): %s", attempt, msg)
            logging.debug(traceback.format_exc())
            if "Call was rejected by callee" in msg or "-2147418111" in msg:
                time.sleep(delay_seconds)
                continue
            raise
    if mail is None:
        raise RuntimeError("Outlook indisponible: CreateItem a échoué après retries.")

    # Forcer le compte d'envoi si demandé
    if from_account:
        try:
            matched = False
            for acct in ns.Session.Accounts:
                smtp = getattr(acct, "SmtpAddress", "")
                if smtp and smtp.lower() == from_account.lower():
                    mail._oleobj_.Invoke(*(64209, 0, 8, 0, acct))  # SendUsingAccount
                    matched = True
                    logging.info(f"[MAIL] SendUsingAccount = {smtp}")
                    break
            if not matched:
                logging.warning(f"[MAIL] Compte '{from_account}' non trouvé dans Outlook. Utilisation du compte par défaut.")
        except Exception as e:
            logging.warning(f"[MAIL] Impossible de forcer le compte '{from_account}': {e}")
            logging.debug(traceback.format_exc())

    # Destinataires & contenu
    mail.To = to_email
    if cc:  mail.CC  = cc
    if bcc: mail.BCC = bcc
    mail.Subject = subject

    final_html = html_body or ""

    try:
        sig_html = ""
        effective_name = ""
        if SIGNATURE_NAME:  # priorité à une signature précise
            sig_html, effective_name = _load_signature_html_by_name(SIGNATURE_NAME)
        elif use_signature:  # sinon signature système
            sig_html = _load_system_signature_html()

        if sig_html:
            # Optionnel : embarquer les images de la signature en CID
            if STRICT_EMBED_SIGNATURE_IMAGES:
                sig_html = _attach_and_inline_signature_images(mail, effective_name or SIGNATURE_NAME, sig_html)
            final_html += sig_html
            logging.debug("[MAIL] Signature ajoutée: %s", (effective_name or SIGNATURE_NAME or "system_default"))
        else:
            logging.debug("[MAIL] Aucune signature chargée (SIGNATURE_NAME=%r, use_system=%r)", SIGNATURE_NAME, use_signature)
    except Exception as e:
        logging.warning(f"[MAIL] Chargement/embarquage signature a échoué: {e}")
        logging.debug(traceback.format_exc())

    mail.HTMLBody = final_html


    # Pièces jointes
    for att in attachments or []:
        try:
            mail.Attachments.Add(str(att))
            logging.debug(f"[MAIL] Attachment ajouté: {att}")
        except Exception as e:
            logging.warning(f"[MAIL] Impossible d'attacher {att}: {e}")
            logging.debug(traceback.format_exc())

    # Sauvegarde avant envoi
    try:
        mail.Save()
        logging.debug("[MAIL] Draft sauvegardé")
    except Exception:
        logging.debug("[MAIL] mail.Save() a échoué (non bloquant)")
        logging.debug(traceback.format_exc())

    # ENVOI avec retries
    for attempt in range(1, max_retries + 1):
        try:
            mail.Send()
            logging.info(f"[MAIL] Send OK -> {to_email}")
            return
        except Exception as e:
            msg = str(e)
            logging.warning("[MAIL] Send échoué (attempt %d): %s", attempt, msg)
            logging.debug(traceback.format_exc())
            if "Call was rejected by callee" in msg or "-2147418111" in msg:
                time.sleep(delay_seconds)
                continue
            raise

    raise RuntimeError(f"[MAIL] Echec envoi à {to_email} après {max_retries} tentatives.")

# --------------------------------------------------------


def main() -> int:
    if not TEMPLATE.exists():
        print(f"[ERREUR] Modèle introuvable: {TEMPLATE}")
        return 1
    if not CSV_FILE.exists():
        print(f"[ERREUR] Fichier CSV introuvable: {CSV_FILE}")
        return 1

    rows = read_rows(CSV_FILE)
    if not rows:
        print("[ERREUR] Aucune ligne valide dans data/entrepreneurs.csv (colonne 'nom' requise).")
        return 1

    print(f"[INFO] Génération pour {len(rows)} entrée(s) d'après le CSV.")

    # Générer DOCX et PDFs
    docx_files, emails, pdf_files = [], [], []
    for i, row in enumerate(rows):
        docx = make_docx_for(row['nom'], i+1)
        docx_files.append(docx)
        emails.append(row.get('email', ''))
        try:
            pdf = convert_docx_to_pdf(docx, emails[-1])
            pdf_files.append(pdf)
        except Exception as e:
            print("[ERREUR] Conversion PDF : docx2pdf nécessite Microsoft Word installé sous Windows.")
            print(f"Détails: {e}")
            print("Alternative: LibreOffice headless -> voir README.md.")
            return 2

    print(f"[OK] {len(docx_files)} DOCX générés -> {OUT_DOCX_DIR}")
    print(f"[OK] {len(pdf_files)} PDF générés -> {OUT_PDF_DIR}")

    # Envoi e-mail (facultatif)
    if SEND_EMAIL:
        sent = 0
        log_env()
        debug_list_outlook_accounts()
        
        for idx, row in enumerate(rows):
            to_email = (row.get("email") or "").strip()
            if not to_email:
                print(f"[INFO] Pas d'email pour la ligne {idx+1}: {row.get('nom')}")
                continue
            subject = SUBJECT_TEMPLATE.format(nom=row.get("nom", ""))
            body_html = BODY_HTML_TEMPLATE.format(nom=row.get("nom", ""))
            pdf_path = pdf_files[idx] if idx < len(pdf_files) else None
            if pdf_path and Path(pdf_path).exists():
                try:
                    send_email_via_outlook(
                        to_email=to_email,
                        subject=subject,
                        html_body=body_html,
                        attachments=[pdf_path],
                        from_account=FROM_ACCOUNT,
                        cc=CC,
                        bcc=BCC,
                        use_signature=USE_SYSTEM_SIGNATURE
                    )
                    sent += 1
                except Exception as e:
                    print(f"[WARN] Echec envoi à {to_email}: {e}")
        print(f"[OK] Emails envoyés: {sent}/{len(rows)}")

    return 0


if __name__ == "__main__":
    sys.exit(main())
